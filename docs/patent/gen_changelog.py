# coding: utf-8
"""
生成DualSoul专利整理工作日志 Word文档
保存到桌面年轮更改说明文件夹
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = '微软雅黑'
OUT_DIR = r'C:\Users\think\Desktop\年轮更改说明'
os.makedirs(OUT_DIR, exist_ok=True)

def sf(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'), FONT)
    rPr.insert(0, rf)

def add_para(doc, text, size=11, bold=False, color=None, indent=0,
             space_before=0, space_after=4, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    r = p.add_run(text)
    sf(r, size=size, bold=bold, color=color)
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        sf(r, size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row[:len(headers)]):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(v))
            sf(r, size=10)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()

# ── 构建文档 ─────────────────────────────────────────────────────────────────

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(1.0)
    s.left_margin = s.right_margin = Inches(1.2)

# 标题页
add_para(doc, 'DualSoul 专利文件整理工作日志', size=20, bold=True,
         color=(31,73,125), space_before=20, space_after=6,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '2026年03月17日', size=12, color=(100,100,100),
         space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '跃创三品文化科技有限公司 · 发明人：程跃', size=11,
         color=(100,100,100), space_after=20, align=WD_ALIGN_PARAGRAPH.CENTER)

# 一、工作概述
add_para(doc, '一、工作概述', size=14, bold=True, color=(31,73,125),
         space_before=10, space_after=6)
add_para(doc, '本次工作对DualSoul项目的全部12份发明专利申请文件进行了系统整理和完善，'
         '确保每份文件符合中国国家知识产权局（CNIPA）发明专利申请的完整性要求，'
         '可直接提交专利代理人或进行CNIPA自助提交。', space_after=4)

add_para(doc, '涉及项目：DualSoul 双身份社交协议', space_after=2)
add_para(doc, 'GitHub：github.com/Chengyue5211/DualSoul', space_after=2)
add_para(doc, '宽限期截止：2026-09-05（专利1-5）/ 2026-09-14（专利6-8）/ 2026-09-15（专利9-12）', space_after=6)

# 二、问题与修复汇总
add_para(doc, '二、本次修复问题汇总', size=14, bold=True, color=(31,73,125),
         space_before=10, space_after=6)

issues = [
    ('申请人/发明人名称错误', '全部12份',
     '原填写GitHub用户名"Chengyue5211"，不符合CNIPA实名要求',
     '申请人→跃创三品文化科技有限公司\n发明人→程跃'),
    ('缺少摘要（六、摘要）', '专利1-12',
     'CNIPA要求说明书必须包含独立摘要章节（150-300字）',
     '为12份专利各自补写150-270字摘要，概括核心技术方案'),
    ('缺少附图说明（四、附图说明）', '专利1-12',
     'CNIPA说明书标准结构要求附图说明章节（无附图时可简略，但章节须存在）',
     '为12份专利各补充3幅附图的文字描述，供代理绘制正式附图'),
    ('缺少具体实施方式（五、具体实施方式）', '专利1-12',
     'CNIPA审查要求"说明书公开充分"，具体实施方式是核心必填章节',
     '为12份专利各补充3个完整实施例，覆盖典型使用场景'),
    ('章节编号混乱', '专利1-5',
     '多次修改后部分章节编号未同步更新，出现重复编号',
     '统一重编为：一至十共10个章节，顺序与CNIPA标准一致'),
]

add_table(doc,
    ['问题类型', '影响范围', '问题描述', '修复方案'],
    issues,
    col_widths=[3.5, 2.0, 5.5, 5.5]
)

# 三、最终文件结构
add_para(doc, '三、最终文件结构（每份专利统一）', size=14, bold=True,
         color=(31,73,125), space_before=10, space_after=6)

structure = [
    ('一、技术领域', '本发明所属技术领域简述', '已有'),
    ('二、背景技术', '现有方案的不足与空白地带', '已有'),
    ('三、发明内容', '技术问题 + 技术方案（含代码示例）+ 有益效果', '已有'),
    ('四、附图说明', '3幅附图的文字描述，供代理绘制（无需发明人画图）', '本次新增'),
    ('五、具体实施方式', '3个完整实施例，满足公开充分性要求', '本次新增'),
    ('六、摘要', '150-270字，概括核心技术方案', '本次新增/完善'),
    ('七、权利要求书', '3-8条权利要求，覆盖独立权项与从属权项', '已有'),
    ('八、在先技术对比', '对比表：在先技术与本发明的区别', '已有'),
    ('九、确权证据', 'Git提交记录、GitHub仓库、许可证信息', '已有'),
    ('十、附件说明', '参考实现源代码、测试用例等附件清单', '已有'),
]

add_table(doc,
    ['章节', '内容说明', '状态'],
    structure,
    col_widths=[3.5, 9.5, 2.5]
)

# 四、12份专利清单
add_para(doc, '四、12份专利完整清单', size=14, bold=True, color=(31,73,125),
         space_before=10, space_after=6)

patents = [
    ('1', '双身份社交图谱架构', '2026-09-05', '核心专利，双节点图谱结构'),
    ('2', 'DISP四模式消息路由协议', '2026-09-05', '核心专利，H2H/H2T/T2H/T2T路由'),
    ('3', '渐进式信任证书系统', '2026-09-05', '核心专利，L0-L4五级TMS评分'),
    ('4', '跨语言人格保真通信方法', '2026-09-05', '核心专利，人格画像注入翻译'),
    ('5', '跨平台分身可移植格式与联邦导入系统', '2026-09-05', 'TPF格式+热冷分层+冲突解决'),
    ('6', '数字分身智能延迟应答方法', '2026-09-14', '三档在线状态+延迟30秒检查'),
    ('7', '对话记忆感知的分身应答系统', '2026-09-14', '滑动窗口6条+角色映射+去重'),
    ('8', 'AI驱动的数字分身视觉形象生成方法', '2026-09-14', '真人照→AI风格化+CSS降级'),
    ('9', '对话叙事记忆与三层聚合系统', '2026-09-15', '时间间隔分段+五维摘要+日级聚合'),
    ('10', '事件驱动分身自主反应引擎', '2026-09-15', '7种事件+6种反应+防抖+生命感'),
    ('11', '分身Agent工具使用系统', '2026-09-16', '关键词检测+两阶段响应+多层降级'),
    ('12', '跨平台Agent互操作API', '2026-09-15', 'API密钥+scope权限+审计日志'),
]

add_table(doc,
    ['序号', '专利名称', '宽限期截止', '核心技术点'],
    patents,
    col_widths=[1.2, 5.5, 3.0, 5.8]
)

# 五、自查结果
add_para(doc, '五、自查结果', size=14, bold=True, color=(31,73,125),
         space_before=10, space_after=6)

checks = [
    ('申请人/发明人名称', '全部12份', 'PASS', '跃创三品文化科技有限公司 / 程跃'),
    ('摘要章节（六、）', '全部12份', 'PASS', '150-270字，符合CNIPA要求'),
    ('附图说明（四、）', '全部12份', 'PASS', '每份3幅附图文字描述'),
    ('具体实施方式（五、）', '全部12份', 'PASS', '每份3个实施例'),
    ('章节编号连续', '全部12份', 'PASS', '一至十，无跳号无重号'),
    ('权利要求书（七、）', '全部12份', 'PASS', '3-8条权项'),
    ('docx文件生成', '全部12份', 'PASS', 'docx/目录下12个文件'),
    ('附图实体文件', '全部12份', 'PENDING', '文字描述已写，图片由代理绘制'),
    ('请求书（申请表）', '全部12份', 'PENDING', '需在CNIPA系统或通过代理单独填写'),
]

add_table(doc,
    ['检查项', '范围', '状态', '备注'],
    checks,
    col_widths=[4.5, 2.0, 2.0, 7.0]
)

# 六、待代理完成事项
add_para(doc, '六、仍需代理完成的工作', size=14, bold=True, color=(31,73,125),
         space_before=10, space_after=6)

todos = [
    '1. 根据四、附图说明的文字描述，绘制每份专利的3幅正式附图（流程图/架构图）',
    '2. 填写CNIPA发明专利申请请求书（申请人、发明人、优先权等信息）',
    '3. 对权利要求书进行专业润色，确保独立权项的保护范围表述准确',
    '4. 确认12份专利的分案策略（是否合并申请或单独申请）',
    '5. 专利1-5宽限期截止2026-09-05，建议2026-05前完成正式提交',
]

for t in todos:
    add_para(doc, t, indent=0.3, space_after=3)

# 七、提交建议
add_para(doc, '七、提交优先级建议', size=14, bold=True, color=(31,73,125),
         space_before=10, space_after=6)

priority = [
    ('第一优先', '专利1、2', '双身份图谱架构 + DISP消息路由', '最核心，技术壁垒最高'),
    ('第一优先', '专利3', '渐进式信任证书系统', '与专利1/2强关联，建议同批提交'),
    ('第二优先', '专利4、5', '跨语言翻译 + 跨平台导入', '功能完整，差异化明显'),
    ('第三优先', '专利6-8', '延迟应答、记忆感知、视觉形象', '宽限期稍晚，可第二批'),
    ('第四优先', '专利9-12', '叙事记忆、事件引擎、工具使用、API', '宽限期最晚，可第三批'),
]

add_table(doc,
    ['优先级', '专利编号', '专利名称', '说明'],
    priority,
    col_widths=[2.0, 2.5, 5.5, 5.5]
)

# 八、本次git提交记录
add_para(doc, '八、本次相关Git提交记录', size=14, bold=True, color=(31,73,125),
         space_before=10, space_after=6)

commits = [
    ('e52e3b6', '2026-03-17', 'feat: 全部12份专利补全附图说明+具体实施方式，结构最终定稿'),
    ('7bb04d2', '2026-03-17', 'fix: 专利6-12补全摘要+章节重编号，全部12份完整定稿'),
    ('6a13f66', '2026-03-17', 'fix: 专利6-12申请人/发明人更正'),
    ('de563bb', '2026-03-17', 'fix: 专利申请人/发明人更正+补全摘要+重新生成docx'),
    ('b9423f9', '2026-03-15', 'ip: 专利9-12交底书 + 确权声明v1.4'),
    ('0ec9a84', '2026-03-14', 'ip: 专利6-8交底书 + 确权声明v1.3'),
    ('d5b3561', '2026-03-14', 'feat: 双服务器并行 — DuckDuckGo搜索恢复'),
]

add_table(doc,
    ['提交Hash', '日期', '提交说明'],
    commits,
    col_widths=[2.5, 2.5, 10.5]
)

# 保存
fname = '2026-03-17_DualSoul专利文件整理完成.docx'
out = os.path.join(OUT_DIR, fname)
doc.save(out)
print('已生成:', out)
