# coding: utf-8
"""
将5份专利 .md 文件转换为格式化 Word 文档
供提交专利代理人或CNIPA自助提交使用
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(__file__), 'docx')
os.makedirs(OUT_DIR, exist_ok=True)

import glob as _glob
PATENTS = sorted([os.path.basename(p) for p in _glob.glob(os.path.join(os.path.dirname(__file__), '专利[0-9]*.md'))])

SRC_DIR = os.path.dirname(__file__)

MYH = '\u5fae\u8f6f\u96c5\u9ed1'

def sf(run, size=11, bold=False, color=None):
    run.font.name = MYH
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'), MYH)
    rPr.insert(0, rf)

def add_para(doc, text, size=11, bold=False, color=None, indent=0, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    sf(r, size=size, bold=bold, color=color)
    return p

def process_md(md_path):
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1.0)
        s.left_margin = s.right_margin = Inches(1.2)

    with open(md_path, encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # H1
        if line.startswith('# ') and not line.startswith('## '):
            add_para(doc, line[2:].strip(), size=16, bold=True, color=(31,73,125), space_before=12, space_after=8)

        # H2
        elif line.startswith('## ') and not line.startswith('### '):
            add_para(doc, line[3:].strip(), size=14, bold=True, color=(31,73,125), space_before=10, space_after=6)

        # H3
        elif line.startswith('### '):
            add_para(doc, line[4:].strip(), size=12, bold=True, color=(54,95,145), space_before=8, space_after=4)

        # H4
        elif line.startswith('#### '):
            add_para(doc, line[5:].strip(), size=11, bold=True, space_before=6, space_after=3)

        # Table
        elif line.startswith('|') and i+1 < len(lines) and lines[i+1].startswith('|---'):
            # collect table rows
            header_row = [c.strip() for c in line.split('|')[1:-1]]
            i += 2  # skip separator
            data_rows = []
            while i < len(lines) and lines[i].startswith('|'):
                data_rows.append([c.strip() for c in lines[i].split('|')[1:-1]])
                i += 1
            # render table
            if header_row and data_rows:
                t = doc.add_table(rows=1+len(data_rows), cols=len(header_row))
                t.style = 'Light Grid Accent 1'
                for ci, h in enumerate(header_row):
                    c = t.rows[0].cells[ci]; c.text = ''
                    r = c.paragraphs[0].add_run(h)
                    sf(r, size=10, bold=True)
                for ri, row in enumerate(data_rows):
                    for ci, v in enumerate(row[:len(header_row)]):
                        # strip markdown bold
                        v = re.sub(r'\*\*(.+?)\*\*', r'\1', v)
                        c = t.rows[ri+1].cells[ci]; c.text = ''
                        r = c.paragraphs[0].add_run(v)
                        sf(r, size=9.5)
                doc.add_paragraph()
            continue

        # Code block
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                r = p.add_run('\n'.join(code_lines))
                r.font.name = 'Courier New'
                r.font.size = Pt(9)

        # HR
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run('\u2015' * 30)
            sf(r, size=9, color=(180,180,180))

        # Numbered list
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(text)
            sf(r, size=10.5)

        # Bullet
        elif line.strip().startswith('- ') or line.strip().startswith('\u00b7 '):
            text = line.strip().lstrip('- ').lstrip('\u00b7 ')
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run('\u2022 ' + text)
            sf(r, size=10.5)

        # Normal paragraph
        elif line.strip():
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line.strip())
            text = re.sub(r'`(.+?)`', r'\1', text)
            add_para(doc, text, size=10.5, space_after=3)

        i += 1

    return doc

total = 0
for fn in PATENTS:
    src = os.path.join(SRC_DIR, fn)
    if not os.path.exists(src):
        print('MISSING: ' + fn)
        continue
    doc = process_md(src)
    out = os.path.join(OUT_DIR, fn.replace('.md', '.docx'))
    doc.save(out)
    print('\u5df2\u751f\u6210: ' + out)
    total += 1

print('\n\u5171\u751f\u6210 %d \u4efd\u4e13\u5229\u6587\u4ef6\uff0c\u4fdd\u5b58\u5230: %s' % (total, OUT_DIR))
