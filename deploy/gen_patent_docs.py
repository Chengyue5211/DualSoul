# -*- coding: utf-8 -*-
import os, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

OUTPUT = os.path.join(os.path.expanduser('~'), 'Desktop', 'DualSoul\u53cc\u8eab\u4efd\u793e\u4ea4\u534f\u8bae', '\u4e13\u5229\u7533\u8bf7\u6750\u6599')
os.makedirs(OUTPUT, exist_ok=True)

FONT = '\u5fae\u8f6f\u96c5\u9ed1'

def set_font(run, name=None, size=12, bold=False, color=None):
    run.font.name = name or FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # For CJK font
    from docx.oxml.ns import qn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name or FONT)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.font.name = FONT
        from docx.oxml.ns import qn
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    return h

def add_para(doc, text, bold=False, size=12, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_multiline_para(doc, text, size=12, indent=False):
    lines = text.split('\n')
    for line in lines:
        if not line.strip():
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, size=size)
        p.paragraph_format.space_after = Pt(3)
        if indent:
            p.paragraph_format.left_indent = Cm(0.5)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10)

    # Set col widths if provided
    if col_widths:
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml
        from lxml import etree
        for row in table.rows:
            for ci, width in enumerate(col_widths):
                row.cells[ci].width = Cm(width)
    return table

# Load data
data_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'patent_data.json')
with open(data_file, 'r', encoding='utf-8-sig') as f:
    raw = json.load(f)

patents = raw['patents']

CN = WD_ALIGN_PARAGRAPH.CENTER
RT = WD_ALIGN_PARAGRAPH.RIGHT

for pat in patents:
    pnum = pat['num']
    fname = pat['filename']
    title = pat['title']
    short_title = pat['short_title']

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Title page
    add_para(doc, '', size=14)
    add_para(doc, '\u53d1\u660e\u4e13\u5229\u7533\u8bf7', bold=True, size=22, align=CN, space_after=16)
    add_para(doc, pnum, bold=True, size=14, align=CN, space_after=8)
    add_para(doc, title, bold=True, size=16, align=CN, space_after=40)
    add_para(doc, '', size=14)

    # Info table
    info_rows = [
        ('\u53d1\u660e\u540d\u79f0', title),
        ('\u53d1\u660e\u4eba', 'Chengyue5211'),
        ('\u7533\u8bf7\u4eba', 'Chengyue5211'),
        ('\u9996\u6b21\u516c\u5f00\u65e5\u671f', pat['first_pub']),
        ('\u516c\u5f00\u5e73\u53f0', 'GitHub (github.com/Chengyue5211/DualSoul) / Gitee\u955c\u50cf'),
        ('\u7533\u8bf7\u5bbd\u9650\u671f\u622a\u6b62', pat['deadline']),
        ('\u6280\u672f\u9886\u57df', pat['tech_field']),
    ]
    add_table(doc, ['\u9879\u76ee', '\u5185\u5bb9'], info_rows, col_widths=[4, 11])

    doc.add_page_break()

    # Section 1
    add_heading(doc, '\u4e00\u3001\u6280\u672f\u9886\u57df', 1)
    add_para(doc, '\u672c\u53d1\u660e\u5c5e\u4e8e' + pat['tech_field'] + '\u9886\u57df\uff0c\u5177\u4f53\u6d89\u53ca' + short_title + '\u76f8\u5173\u7684\u7cfb\u7edf\u53ca\u65b9\u6cd5\u3002', size=12)

    # Section 2
    add_heading(doc, '\u4e8c\u3001\u80cc\u666f\u6280\u672f', 1)
    add_multiline_para(doc, pat['background'])

    # Section 3
    add_heading(doc, '\u4e09\u3001\u53d1\u660e\u5185\u5bb9', 1)

    add_heading(doc, '3.1 \u6280\u672f\u95ee\u9898', 2)
    add_para(doc, pat['problem'], size=12)

    add_heading(doc, '3.2 \u6280\u672f\u65b9\u6848', 2)
    for subsec in pat['solution_sections']:
        add_heading(doc, subsec['title'], 3)
        add_multiline_para(doc, subsec['content'])

    add_heading(doc, '3.3 \u6709\u76ca\u6548\u679c', 2)
    add_multiline_para(doc, pat['benefits'])

    # Section 4 - Claims
    add_heading(doc, '\u56db\u3001\u6743\u5229\u8981\u6c42\u4e66', 1)
    for i, claim in enumerate(pat['claims'], 1):
        p = doc.add_paragraph()
        run = p.add_run(str(i) + '. ' + claim)
        set_font(run, size=12, bold=(i == 1))
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.3)

    # Section 5 - Prior art table
    add_heading(doc, '\u4e94\u3001\u5728\u5148\u6280\u672f\u5bf9\u6bd4', 1)
    add_table(doc, ['\u5728\u5148\u6280\u672f', '\u4e0e\u672c\u53d1\u660e\u7684\u533a\u522b'], pat['prior_art'], col_widths=[5, 10])
    add_para(doc, '')
    add_para(doc, '\u672c\u53d1\u660e\u7684\u72ec\u521b\u6027\uff1a' + pat['innovation'], bold=True, size=12)

    # Section 6 - Evidence
    add_heading(doc, '\u516d\u3001\u786e\u6743\u8bc1\u636e', 1)
    add_para(doc, '\u672c\u53d1\u660e\u7684\u9996\u6b21\u516c\u5f00\u8bb0\u5f55\u5982\u4e0b\uff1a', size=12)
    ev_rows = [
        ('Git\u63d0\u4ea4\u8bb0\u5f55', pat['git_hash']),
        ('GitHub\u4ed3\u5e93', 'github.com/Chengyue5211/DualSoul'),
        ('Gitee\u955c\u50cf', 'gitee.com/chengyue5211/DualSoul'),
        ('\u8bb8\u53ef\u8bc1', 'AGPL-3.0-or-later'),
        ('\u5173\u8054\u9879\u76ee', '\u8dc3\u521b\u5e74\u8f6e (ycnianlun.com) \u2014 \u53cc\u8eab\u4efd\u793e\u4ea4\u6a21\u5757\u5df2\u96c6\u6210'),
    ]
    add_table(doc, ['\u8bc1\u636e\u7c7b\u578b', '\u5185\u5bb9'], ev_rows, col_widths=[4, 11])

    # Section 7 - Annexes
    add_heading(doc, '\u4e03\u3001\u9644\u4ef6\u8bf4\u660e', 1)
    annexes = [
        '1. GitHub\u4ed3\u5e93\u5b8c\u6574\u63d0\u4ea4\u5386\u53f2\uff08github.com/Chengyue5211/DualSoul\uff09',
        '2. Gitee\u955c\u50cf\u4ed3\u5e93\u63d0\u4ea4\u5386\u53f2\uff08gitee.com/chengyue5211/DualSoul\uff09',
        '3. \u4e13\u5229\u6280\u672f\u4ea4\u5e95\u4e66\uff08docs/PATENT_DISCLOSURE.md\uff09',
        '4. \u77e5\u8bc6\u4ea7\u6743\u786e\u6743\u58f0\u660e\uff08docs/IP_CONFIRMATION_v1.1.md\uff09',
        '5. \u767d\u76ae\u4e66\uff08docs/whitepaper.md v1.1\uff09',
        '6. \u534f\u8bae\u89c4\u8303\uff08docs/protocol.md v1.1\uff09',
        '7. \u53c2\u8003\u5b9e\u73b0\u6e90\u4ee3\u7801\uff08dualsoul/\uff09',
        '8. \u81ea\u52a8\u5316\u6d4b\u8bd5\uff08tests/\uff0c35\u4e2a\u6d4b\u8bd5\u7528\u4f8b\uff09',
    ]
    for a in annexes:
        add_para(doc, a, size=12, space_after=4)

    add_para(doc, '')
    add_para(doc, '\u2014' * 30)
    now_str = datetime.now().strftime('%Y\u5e74%m\u6708%d\u65e5')
    add_para(doc, '\u53d1\u660e\u4eba\uff1aChengyue5211  |\u65e5\u671f\uff1a' + now_str, size=10, align=RT)
    add_para(doc, '\u672c\u6587\u6863\u4ec5\u7528\u4e8e\u4e13\u5229\u7533\u8bf7\u51c6\u5907\uff0c\u4e0d\u6784\u6210\u516c\u5f00\u7684\u6280\u672f\u62ab\u9732', size=10, align=RT)

    out_path = os.path.join(OUTPUT, fname)
    doc.save(out_path)
    print('OK: ' + out_path)

# ---- Guide document ----
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

add_para(doc, 'DualSoul \u4e13\u5229\u7533\u8bf7\u6750\u6599\u6c47\u603b\u6307\u5357', bold=True, size=20, align=CN, space_after=12)
now_str = datetime.now().strftime('%Y\u5e74%m\u6708%d\u65e5')
add_para(doc, '\u751f\u6210\u65e5\u671f\uff1a' + now_str + '  |\u53d1\u660e\u4eba\uff1aChengyue5211', size=12, align=CN, space_after=30)

add_heading(doc, '\u4e00\u3001\u6750\u6599\u6e05\u5355', 1)
mat_rows = [
    ('\u4e13\u52491', '\u53cc\u8eab\u4efd\u793e\u4ea4\u56fe\u8c31\u67b6\u6784', '2026-03-05', '2026-09-05', '\u9ad8'),
    ('\u4e13\u52492', 'DISP\u56db\u6a21\u5f0f\u6d88\u606f\u8def\u7531\u534f\u8bae', '2026-03-05', '2026-09-05', '\u9ad8'),
    ('\u4e13\u52493', '\u6e10\u8fdb\u5f0f\u4fe1\u4efb\u8bc1\u4e66\u7cfb\u7edf', '2026-03-05', '2026-09-05', '\u4e2d'),
    ('\u4e13\u52494', '\u8de8\u8bed\u8a00\u4eba\u683c\u4fdd\u771f\u901a\u4fe1\u65b9\u6cd5', '2026-03-13', '2026-09-13', '\u4e2d'),
]
add_table(doc, ['\u7f16\u53f7', '\u6838\u5fc3\u521b\u65b0', '\u9996\u6b21\u516c\u5f00', '\u5e94\u8bf7\u622a\u6b62', '\u4f18\u5148\u7ea7'], mat_rows)

add_heading(doc, '\u4e8c\u3001\u65f6\u95f4\u8282\u70b9', 1)
timeline_rows = [
    ('2026-03-05', '\u4e13\u52491\u30012\u30013\u9996\u6b21\u516c\u5f00\uff08GitHub\u63d0\u4ea4\u8bb0\u5f55\u4e3a\u8bc1\uff09'),
    ('2026-03-13', '\u4e13\u52494\u9996\u6b21\u516c\u5f00\uff08\u5f53\u65e5\u65e5\u671f\uff09'),
    ('2026-06-30', '\u5efa\u8bae\u63d0\u4ea4\u622a\u6b62\uff08\u7ed9\u4ee3\u7406\u673a\u6784\u751f\u6210\u6b63\u5f0f\u7533\u8bf7\u6587\u4ef6\u7559\u52a1\u8db3\u591f\u65f6\u95f4\uff09'),
    ('2026-09-05', '\u4e13\u52491\u30012\u30013\u5168\u56fd\u4eba\u5927\u5e38\u59d4\u4f1a\u5c0f\u5e74\u5ba4\u671f\u622a\u6b62\uff08\u9996\u6b21\u516c\u5f00\u540e6\u4e2a\u6708\u5bbd\u9650\u671f\uff09'),
    ('2026-09-13', '\u4e13\u52494\u5bbd\u9650\u671f\u622a\u6b62'),
]
add_table(doc, ['\u65e5\u671f', '\u4e8b\u9879'], timeline_rows)

add_heading(doc, '\u4e09\u3001\u7533\u8bf7\u5efa\u8bae', 1)
suggestions = [
    '1. \u59d4\u6258\u4ee3\u7406\u673a\u6784\uff1a\u5efa\u8bae\u59d4\u6258\u6709AI/\u8f6f\u4ef6\u7c7b\u4e13\u5229\u7533\u8bf7\u7ecf\u9a8c\u7684\u4ee3\u7406\u673a\u6784\u5904\u7406\u6b63\u5f0f\u7533\u8bf7\u6587\u4ef6\uff08\u6743\u5229\u8981\u6c42\u4e66\u3001\u8bf4\u660e\u4e66\u3001\u6458\u8981\uff09',
    '2. \u8d39\u7528\u9884\u4f30\uff1a\u53d1\u660e\u4e13\u5229\u6bcf\u9879\u5b98\u65b9\u8d39\u7528\u7ea69000\u5143\uff08\u5927\u81f3250\u5143+\u7533\u8bf7\u8d39\u950090\u5143+\u8bf4\u660e\u4e66\u4ee3\u4ee3\u8d39\u8d39\uff09\uff0c\u52a0\u4e0a\u4ee3\u7406\u673a\u6784\u670d\u52a1\u8d39\u7ea63000\u22128000\u5143/\u9879\uff0c\u56db\u9879\u5408\u8ba112000\u221232000\u5143',
    '3. \u4f18\u5148\u7ea7\u5efa\u8bae\uff1a\u5982\u9884\u7b97\u6709\u9650\uff0c\u4f18\u5148\u7533\u8bf7\u4e13\u52491\uff08\u53cc\u8eab\u4efd\u5ea6\u56fe\u8c31\u67b6\u6784\uff0c\u4e3b\u67b6\u6784\u7c7b\uff09\u548c\u4e13\u52492\uff08DISP\u534f\u8bae\uff0c\u534f\u8bae\u7c7b\uff09\uff0c\u8fd9\u4e24\u9879\u4fdd\u62a4\u8303\u56f4\u6700\u5e7f',
    '4. PCT\u56fd\u9645\u7533\u8bf7\uff1a\u82e5\u6709\u56fd\u9645\u5e02\u573a\u8ba1\u5212\uff0c\u53ef\u5728\u4e2d\u56fd\u9996\u6b21\u7533\u8bf7\u540e12\u4e2a\u6708\u5185\u63d0\u4ea4PCT\u7533\u8bf7\uff0c\u5c45\u540e\u8fdb\u5165\u7279\u5b9a\u56fd\u5bb6/\u5730\u533a\u7684\u56fd\u5185\u9636\u6bb5',
    '5. \u8f6f\u4ef6\u8457\u4f5c\u6743\uff1a\u5df2\u5907\u8f6f\u4ef6\u8457\u4f5c\u6743\u767b\u8bb0\u6750\u6599\uff0c\u53ef\u540c\u6b65\u7533\u8bf7\uff0c\u8865\u5145\u4fdd\u62a4',
]
for s in suggestions:
    add_para(doc, s, size=12, space_after=8)

add_heading(doc, '\u56db\u3001\u652f\u64c1\u6750\u6599\u6e05\u5355', 1)
support_rows = [
    ('PATENT_DISCLOSURE.md', '\u6280\u672f\u4ea4\u5e95\u4e66\uff08\u672c\u7ec4\u6587\u6863\u7684\u539f\u59cb\u8d44\u6599\uff09', '\u5df2\u5907'),
    ('IP_CONFIRMATION_v1.1.md', '\u77e5\u8bc6\u4ea7\u6743\u786e\u6743\u58f0\u660e', '\u5df2\u5907'),
    ('whitepaper.md', 'DualSoul\u767d\u76ae\u4e66v1.1', '\u5df2\u5907'),
    ('protocol.md', 'DISP\u534f\u8bae\u89c4\u8303v1.1', '\u5df2\u5947'),
    ('dualsoul/\u6e90\u4ee3\u7801', '\u53c2\u8003\u5b9e\u73b0\uff08Python\uff09', '\u5df2\u5907'),
    ('tests/\u6d4b\u8bd5\u7528\u4f8b', '35\u4e2apytest\u6d4b\u8bd5', '\u5df2\u5907'),
    ('GitHub\u63d0\u4ea4\u5386\u53f2', '\u9996\u6b21\u516c\u5f00\u65e5\u671f\u8bc1\u636e', '\u5df2\u5907'),
    ('Gitee\u955c\u50cf', '\u56fd\u5185\u5907\u4efd\uff0c\u53ef\u4f5c\u4e8c\u91cd\u8bc1\u636e', '\u5df2\u5907'),
]
add_table(doc, ['\u6587\u4ef6/\u8d44\u6e90', '\u7528\u9014', '\u72b6\u6001'], support_rows)

add_para(doc, '')
add_para(doc, '\u2014' * 30)
add_para(doc, '\u672c\u6307\u5357\u7531Claude Code\u81ea\u52a8\u751f\u6210 | \u4ec5\u4f9b\u4e13\u5229\u7533\u8bf7\u51c6\u5907\u53c2\u8003', size=10, align=RT)

guide_path = os.path.join(OUTPUT, '00_\u4e13\u5229\u7533\u8bf7\u6307\u5357.docx')
doc.save(guide_path)
print('OK: ' + guide_path)
print('ALL DONE - 5 documents generated successfully!')
